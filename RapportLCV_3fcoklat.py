import streamlit as st
import pandas as pd
import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
import io
import PyPDF2
from PIL import Image
import pytesseract
import time
import hashlib

# Set path Tesseract untuk Windows
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Konfigurasi API DeepSeek (AMAN melalui secrets)
DEEPSEEK_API_KEY = st.secrets["deepseek"]["api_key"]
DEEPSEEK_MODEL = "deepseek-r1-0528"  # Model terbaik untuk keseimbangan kualitas-biaya (Mei 2025)

# Konfigurasi halaman
st.set_page_config(
    page_title="Rapport Writer Assistance",
    page_icon="📊",
    layout="wide"
)

# Fungsi untuk normalisasi nama HSH
def normalize_hsh(hsh_name):
    if pd.isna(hsh_name):
        return ""
    normalized = str(hsh_name).strip().upper()
    normalized = ' '.join(normalized.split())
    return normalized

def find_matching_hsh(target_hsh, hsh_list):
    target_normalized = normalize_hsh(target_hsh)
    for hsh in hsh_list:
        if normalize_hsh(hsh) == target_normalized:
            return hsh
    for hsh in hsh_list:
        normalized = normalize_hsh(hsh)
        if target_normalized in normalized or normalized in target_normalized:
            return hsh
    return None

@st.cache_data
def load_excel_files():
    try:
        skor_total = pd.read_excel('documents/SKOR_TOTAL_ALL.xlsx', sheet_name='SKOR TOTAL_ALL')
        skor_survei = pd.read_excel('documents/Skor_SURVEI_ALL.xlsx', sheet_name='Skor_SURVEI_ALL_FUNGSI')
        skor_benchmark_evidence = pd.read_excel('documents/Skor_benchmark.xlsx', sheet_name='Evidence')
        skor_benchmark_survei = pd.read_excel('documents/Skor_benchmark.xlsx', sheet_name='Survei')
        
        if 'HSH' in skor_total.columns:
            skor_total['HSH_normalized'] = skor_total['HSH'].apply(normalize_hsh)
        if 'HSH' in skor_survei.columns:
            skor_survei['HSH_normalized'] = skor_survei['HSH'].apply(normalize_hsh)
        
        skor_benchmark_evidence['HSH_normalized'] = skor_benchmark_evidence.iloc[:, 0].apply(normalize_hsh)
        skor_benchmark_survei['HSH_normalized'] = skor_benchmark_survei.iloc[:, 0].apply(normalize_hsh)
        
        return skor_total, skor_survei, skor_benchmark_evidence, skor_benchmark_survei
    except Exception as e:
        st.error(f"Error loading Excel files: {str(e)}")
        st.info("Pastikan folder 'documents' ada dan berisi file: SKOR_TOTAL_ALL.xlsx, Skor_SURVEI_ALL.xlsx, dan Skor_benchmark.xlsx")
        return None, None, None, None

def extract_text_from_pdf(pdf_file):
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

def extract_text_from_image(image_file):
    try:
        image = Image.open(image_file)
        text = pytesseract.image_to_string(image, lang='ind+eng')
        return text
    except Exception as e:
        return f"Error reading image: {str(e)}"

def read_uploaded_file(uploaded_file):
    if uploaded_file is None:
        return None
    
    file_extension = uploaded_file.name.split('.')[-1].lower()
    
    try:
        if file_extension in ['xlsx', 'xls']:
            df = pd.read_excel(uploaded_file)
            return df.to_string()
        elif file_extension == 'pdf':
            return extract_text_from_pdf(uploaded_file)
        elif file_extension in ['png', 'jpg', 'jpeg']:
            return extract_text_from_image(uploaded_file)
        else:
            return "Format file tidak didukung"
    except Exception as e:
        return f"Error reading file: {str(e)}"

# Fungsi untuk memanggil DeepSeek API dengan retry mechanism
def call_deepseek(prompt, max_tokens=1500, max_retries=2):
    """Memanggil DeepSeek API dengan retry mechanism dan parameter optimasi"""
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }
    
    # System prompt yang diperbaiki dengan fokus pada reasoning
    system_prompt = """Anda adalah konsultan senior budaya kerja perusahaan yang berpengalaman dengan pendekatan apresiatif dan profesional. 

**INSTRUKSI KHUSUS:**
1. BERIKAN REASONING LENGKAP sebelum kesimpulan akhir
2. Gunakan data yang disediakan secara spesifik dan terukur
3. Setiap rekomendasi harus memiliki DASAR ANALISIS yang jelas
4. Jelaskan LOGIKA di balik setiap poin yang Anda sampaikan
5. FOKUS pada aspek PERILAKU (behavior): perubahan mindset, kolaborasi, komunikasi, kepemimpinan, keterlibatan, partisipasi

TONE & GAYA KOMUNIKASI:
- Gunakan bahasa yang apresiatif, menghargai usaha yang telah dilakukan
- Profesional namun hangat dan mendukung
- Fokus pada kekuatan (strength-based approach) sebelum memberikan saran perbaikan
- Hindari kata-kata negatif atau menghakimi
- Gunakan frasa seperti "telah menunjukkan komitmen yang baik", "dapat lebih dioptimalkan", "peluang untuk pengembangan lebih lanjut"
- Berikan apresiasi spesifik terhadap pencapaian yang ada

FORMAT OUTPUT:
- Mulai dengan apresiasi umum
- "Hal yang Sudah Baik" harus spesifik dan menghargai pencapaian
- "Hal yang Dapat Diperbaiki" disampaikan sebagai peluang pengembangan, bukan kritik
- Setiap poin harus memiliki REASONING yang jelas"""

    data = {
        "model": DEEPSEEK_MODEL,
        "messages": [
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0.3,  # Lebih rendah untuk konsistensi dan profesionalisme
        "top_p": 0.9,
        "frequency_penalty": 0.1,
        "presence_penalty": 0.1,
        "max_tokens": max_tokens
    }
    
    for attempt in range(max_retries):
        try:
            response = requests.post("https://api.deepseek.com/v1/chat/completions", headers=headers, json=data)
            if response.status_code == 200:
                result = response.json()
                return result['choices'][0]['message']['content']
            elif response.status_code == 429:  # Rate limit
                if attempt < max_retries - 1:
                    time.sleep(2 ** attempt)  # Exponential backoff
                    continue
                return f"Error: Rate limit exceeded. Silakan coba lagi beberapa saat."
            else:
                return f"Error calling DeepSeek API: {response.status_code} - {response.text}"
        except Exception as e:
            if attempt == max_retries - 1:
                return f"DeepSeek API failed after {max_retries} attempts: {str(e)}"
            time.sleep(1)
    
    return "Gagal menghubungi DeepSeek API setelah beberapa percobaan."

# Fungsi helper untuk caching berdasarkan hash content
def get_content_hash(content):
    if content is None:
        return "empty"
    return hashlib.md5(str(content).encode()).hexdigest()

# Fungsi analisis dengan caching optimasi
@st.cache_data(ttl=3600)  # Cache selama 1 jam
def cached_analyze_strategi_budaya(pcb_hash, selected_hsh, selected_fungsi):
    # Fungsi ini hanya untuk caching, analisis sebenarnya dilakukan di luar
    return None

def analyze_strategi_budaya(pcb_content, selected_hsh, selected_fungsi):
    # Generate hash untuk caching
    pcb_hash = get_content_hash(pcb_content)
    
    # Cek apakah hasil sudah di-cache
    cached_result = cached_analyze_strategi_budaya(pcb_hash, selected_hsh, selected_fungsi)
    if cached_result is not None:
        return cached_result
    
    prompt = f"""
**KONTEKS:** Analisis strategi budaya kerja untuk fungsi {selected_fungsi} di HSH {selected_hsh}

**DATA PCB YANG DISEDIAKAN:**
{pcb_content}

**INSTRUKSI ANALISIS:**
1. IDENTIFIKASI pola dan tren dalam strategi budaya yang diajukan
2. ANALISIS kesesuaian dengan metode SMART dan logika bisnis
3. BERIKAN REASONING untuk setiap poin analisis
4. FOKUS utama pada aspek PERILAKU (bukan teknis operasional)

**ASPEK YANG DINILAI:**
- Penggunaan metode SMART pada Goals/Business Initiatives
- Kerunutan logis dari identifikasi kendala ke Business Initiatives
- Kelengkapan PCB dalam menggambarkan strategi budaya
- Fokus pada perubahan perilaku (mindset, kolaborasi, komunikasi, dll)

**FORMAT OUTPUT YANG DIHARAPKAN:**

**Apresiasi Umum:**
[Berikan apresiasi terhadap upaya dan komitmen yang telah ditunjukkan, fokus pada aspek positif strategi budaya]

**Hal yang Sudah Baik:**
- [Poin spesifik 1] - **Reasoning:** [Penjelasan mengapa ini baik dan dampaknya pada perilaku]
- [Poin spesifik 2] - **Reasoning:** [Penjelasan mengapa ini baik dan dampaknya pada perilaku]

**Peluang Pengembangan Lebih Lanjut:**
- [Saran 1] - **Reasoning:** [Penjelasan logis dan dampak perilaku yang diharapkan]
- [Saran 2] - **Reasoning:** [Penjelasan logis dan dampak perilaku yang diharapkan]

**Rekomendasi Prioritas:**
[Berikan 1-2 rekomendasi utama dengan reasoning yang kuat]
"""
    return call_deepseek(prompt, max_tokens=1500)

@st.cache_data(ttl=3600)
def cached_analyze_program_budaya(pcb_hash, selected_hsh, selected_fungsi):
    return None

def analyze_program_budaya(pcb_content, selected_hsh, selected_fungsi):
    pcb_hash = get_content_hash(pcb_content)
    cached_result = cached_analyze_program_budaya(pcb_hash, selected_hsh, selected_fungsi)
    if cached_result is not None:
        return cached_result
    
    prompt = f"""
**KONTEKS:** Analisis Program Budaya untuk fungsi {selected_fungsi} di HSH {selected_hsh}

**DATA PROGRAM YANG DISEDIAKAN:**
{pcb_content}

**INSTRUKSI ANALISIS:**
1. EVALUASI setiap program (One Hour Meeting, ONE Action, ONE KOLAB) secara terpisah
2. ANALISIS dampak perilaku dari setiap program
3. BERIKAN REASONING untuk setiap evaluasi dan saran
4. FOKUS pada keterlibatan pekerja dan perubahan perilaku nyata

**ASPEK YANG DINILAI PER PROGRAM:**
- Kesesuaian judul dengan tujuan perubahan perilaku
- Kualitas deliverables dalam mendorong perubahan perilaku
- Kontribusi program terhadap pencapaian Goals/Business Initiatives
- Tingkat keterlibatan dan partisipasi pekerja

**FORMAT OUTPUT YANG DIHARAPKAN:**

**Apresiasi Umum:**
[Apresiasi terhadap desain dan implementasi program budaya, soroti komitmen tim]

**Hal yang Sudah Baik:**
- [Program spesifik 1] - **Reasoning:** [Penjelasan dampak perilaku positif yang terukur]
- [Program spesifik 2] - **Reasoning:** [Penjelasan dampak perilaku positif yang terukur]

**Peluang Pengembangan Lebih Lanjut:**
- [Program 1] - **Reasoning:** [Penjelasan logis untuk optimalisasi dampak perilaku]
- [Program 2] - **Reasoning:** [Penjelasan logis untuk optimalisasi dampak perilaku]

**Rekomendasi Implementasi:**
[Berikan 1-2 saran implementasi spesifik dengan reasoning yang kuat]
"""
    return call_deepseek(prompt, max_tokens=1500)

@st.cache_data(ttl=3600)
def cached_analyze_impact(impact_hash, selected_hsh, selected_fungsi):
    return None

def analyze_impact(impact_content, selected_hsh, selected_fungsi):
    if impact_content is None:
        return "Analisis impact tidak dapat dilakukan karena tidak ada file impact to business yang di upload"
    
    impact_hash = get_content_hash(impact_content)
    cached_result = cached_analyze_impact(impact_hash, selected_hsh, selected_fungsi)
    if cached_result is not None:
        return cached_result
    
    prompt = f"""
**KONTEKS:** Analisis Impact to Business untuk fungsi {selected_fungsi} di HSH {selected_hsh}

**DATA IMPACT YANG DISEDIAKAN:**
{impact_content}

**INSTRUKSI ANALISIS:**
1. IDENTIFIKASI perubahan perilaku yang terukur dari data impact
2. ANALISIS hubungan sebab-akibat antara perubahan perilaku dan dampak bisnis
3. BERIKAN REASONING untuk setiap temuan dan rekomendasi
4. FOKUS pada aspek perilaku yang berdampak pada kinerja bisnis

**ASPEK PERILAKU YANG DINILAI:**
- Peningkatan kolaborasi dan kerja sama tim
- Perbaikan komunikasi dan koordinasi
- Perubahan mindset dan budaya kerja
- Peningkatan kepemimpinan dan ownership
- Peningkatan keterlibatan dan motivasi pekerja
- Penerapan nilai-nilai AKHLAK dalam praktik kerja

**FORMAT OUTPUT YANG DIHARAPKAN:**

**Apresiasi Pencapaian:**
[Apresiasi terhadap dampak positif yang telah dicapai, soroti perubahan perilaku yang signifikan]

**Hal yang Sudah Baik:**
- [Perubahan perilaku 1] - **Reasoning:** [Penjelasan dampak bisnis yang terukur]
- [Perubahan perilaku 2] - **Reasoning:** [Penjelasan dampak bisnis yang terukur]

**Peluang Pengembangan Lebih Lanjut:**
- [Area 1] - **Reasoning:** [Penjelasan logis untuk memperkuat dampak perilaku]
- [Area 2] - **Reasoning:** [Penjelasan logis untuk memperkuat dampak perilaku]

**Strategi Pengembangan:**
[Berikan 1-2 strategi spesifik untuk meningkatkan dampak perilaku ke bisnis]
"""
    return call_deepseek(prompt, max_tokens=1500)

@st.cache_data(ttl=3600)
def cached_analyze_evidence_comparison(hsh_hash, fungsi_hash):
    return None

def analyze_evidence_comparison(skor_total, skor_benchmark_evidence, selected_hsh, selected_fungsi):
    try:
        # Generate hash untuk caching
        cache_hash = get_content_hash(f"{selected_hsh}_{selected_fungsi}")
        cached_result = cached_analyze_evidence_comparison(get_content_hash(selected_hsh), get_content_hash(selected_fungsi))
        if cached_result is not None:
            return cached_result
        
        fungsi_data = skor_total[skor_total['Fungsi'] == selected_fungsi]
        if fungsi_data.empty:
            return "Data fungsi tidak ditemukan dalam file SKOR_TOTAL_ALL"
        
        fungsi_hsh = fungsi_data.iloc[0]['HSH'] if 'HSH' in fungsi_data.columns else selected_hsh
        fungsi_hsh_normalized = normalize_hsh(fungsi_hsh)
        
        benchmark_data = skor_benchmark_evidence[
            skor_benchmark_evidence['HSH_normalized'] == fungsi_hsh_normalized
        ]
        
        if benchmark_data.empty:
            st.warning(f"⚠️ HSH '{fungsi_hsh}' tidak ditemukan exact match di benchmark. Mencoba fuzzy matching...")
            match_found = False
            for idx, row in skor_benchmark_evidence.iterrows():
                benchmark_hsh_norm = row['HSH_normalized']
                if fungsi_hsh_normalized in benchmark_hsh_norm or benchmark_hsh_norm in fungsi_hsh_normalized:
                    benchmark_data = skor_benchmark_evidence.iloc[[idx]]
                    st.info(f"✓ Ditemukan match: '{row.iloc[0]}' untuk HSH '{fungsi_hsh}'")
                    match_found = True
                    break
            
            if not match_found:
                st.warning(f"⚠️ Data benchmark untuk HSH '{fungsi_hsh}' tidak ditemukan. Menggunakan benchmark 'Pertamina Group' sebagai referensi.")
                benchmark_data = skor_benchmark_evidence[
                    skor_benchmark_evidence['HSH_normalized'].str.contains('PERTAMINA GROUP', na=False)
                ]
                if benchmark_data.empty:
                    benchmark_data = skor_benchmark_evidence.iloc[[0]]
                    st.info(f"Menggunakan benchmark: '{benchmark_data.iloc[0, 0]}'")
        
        kolom_names = ['Strategi Budaya', 'Monitoring & Evaluasi', 'Sosialisasi & Partisipasi', 
                       'Pelaporan Bulanan', 'Apresiasi Pelanggan', 'Pemahaman Program', 
                       'Reward & Consequences', 'SK AoC', 'Impact to Business']
        fungsi_values = {}
        for i, name in enumerate(kolom_names):
            col_idx = 3 + i
            if col_idx < len(fungsi_data.columns):
                fungsi_values[name] = fungsi_data.iloc[0, col_idx]
            else:
                fungsi_values[name] = 'N/A'
        
        benchmark_values = {}
        for i, name in enumerate(kolom_names):
            col_idx = 1 + i
            if col_idx < len(benchmark_data.columns):
                benchmark_values[name] = benchmark_data.iloc[0, col_idx]
            else:
                benchmark_values[name] = 'N/A'
        
        differences = {}
        for name in kolom_names:
            if name in fungsi_values and name in benchmark_values:
                try:
                    fungsi_val = float(str(fungsi_values[name]).replace(',', '.'))
                    benchmark_val = float(str(benchmark_values[name]).replace(',', '.'))
                    diff = fungsi_val - benchmark_val
                    differences[name] = diff
                except (ValueError, TypeError):
                    differences[name] = 'N/A'
        
        benchmark_hsh_display = benchmark_data.iloc[0, 0] if not benchmark_data.empty else "Benchmark tidak tersedia"
        
        comparison_text = f"""
PERBANDINGAN EVIDENCE

Fungsi: {selected_fungsi}
HSH Fungsi: {fungsi_hsh}
HSH Benchmark: {benchmark_hsh_display}

=== DATA FUNGSI ===
"""
        for name, value in fungsi_values.items():
            comparison_text += f"- {name}: {value}\n"
        
        comparison_text += f"""
=== BENCHMARK ({benchmark_hsh_display}) ===
"""
        for name, value in benchmark_values.items():
            comparison_text += f"- {name}: {value}\n"
        
        comparison_text += f"""
=== SELISIH (Fungsi - Benchmark) ===
"""
        for name, diff in differences.items():
            if diff != 'N/A':
                try:
                    diff_float = float(diff)
                    status = "✓ LEBIH BAIK" if diff_float > 0 else "⚠ PELUANG PENGEMBANGAN" if diff_float < 0 else "= SESUAI"
                    comparison_text += f"- {name}: {diff_float:+.2f} {status}\n"
                except:
                    comparison_text += f"- {name}: {diff} (Data tidak valid)\n"
            else:
                comparison_text += f"- {name}: {diff}\n"
        
        comparison_text += """
Catatan:
- Nilai positif (+) = Fungsi LEBIH BAIK dari benchmark
- Nilai negatif (-) = Fungsi memiliki PELUANG PENGEMBANGAN
"""
        
        prompt = f"""
**KONTEKS:** Analisis perbandingan Evidence untuk fungsi {selected_fungsi} di HSH {selected_hsh}

**DATA PERBANDINGAN YANG DISEDIAKAN:**
{comparison_text}

**INSTRUKSI ANALISIS:**
1. ANALISIS perbedaan skor pada setiap aspek dengan benchmark
2. BERIKAN REASONING untuk setiap temuan
3. FOKUS pada aspek perilaku dalam implementasi budaya kerja
4. IDENTIFIKASI pola dan area prioritas pengembangan

**ASPEK YANG DINILAI:**
- Strategi Budaya dan implementasinya
- Monitoring & Evaluasi oleh AoC dan Pimpinan
- Sosialisasi & Partisipasi dalam program budaya
- Sistem pelaporan dan apresiasi
- Pemahaman program dan sistem reward
- Impact to Business dari program budaya

**FORMAT OUTPUT YANG DIHARAPKAN:**

**Apresiasi Pencapaian:**
[Apresiasi terhadap area yang sudah di atas atau sesuai benchmark, soroti komitmen dan konsistensi]

**Hal yang Sudah Baik:**
- [Area spesifik 1] - **Reasoning:** [Penjelasan mengapa area ini unggul dan dampak perilakunya]
- [Area spesifik 2] - **Reasoning:** [Penjelasan mengapa area ini unggul dan dampak perilakunya]

**Peluang Pengembangan Lebih Lanjut:**
- [Area 1] - **Reasoning:** [Penjelasan logis dan rekomendasi konkret berbasis perilaku]
- [Area 2] - **Reasoning:** [Penjelasan logis dan rekomendasi konkret berbasis perilaku]

**Prioritas Aksi:**
[Berikan 1-2 prioritas aksi spesifik dengan reasoning yang kuat]
"""
        return call_deepseek(prompt, max_tokens=1500)
    except Exception as e:
        error_msg = f"Error dalam analisis evidence: {str(e)}\n\nDetail error: {e.__class__.__name__}"
        st.error(error_msg)
        return error_msg

@st.cache_data(ttl=3600)
def cached_analyze_survei_comparison(hsh_hash, fungsi_hash):
    return None

def analyze_survei_comparison(skor_survei, skor_benchmark_survei, selected_hsh, selected_fungsi):
    try:
        # Generate hash untuk caching
        cache_hash = get_content_hash(f"{selected_hsh}_{selected_fungsi}")
        cached_result = cached_analyze_survei_comparison(get_content_hash(selected_hsh), get_content_hash(selected_fungsi))
        if cached_result is not None:
            return cached_result
        
        fungsi_data = skor_survei[skor_survei['Fungsi'] == selected_fungsi]
        if fungsi_data.empty:
            return "Data survei fungsi tidak ditemukan dalam file Skor_SURVEI_ALL"
        
        fungsi_hsh = fungsi_data.iloc[0]['HSH'] if 'HSH' in fungsi_data.columns else selected_hsh
        fungsi_hsh_normalized = normalize_hsh(fungsi_hsh)
        
        benchmark_data = skor_benchmark_survei[
            skor_benchmark_survei['HSH_normalized'] == fungsi_hsh_normalized
        ]
        
        if benchmark_data.empty:
            st.warning(f"⚠️ HSH '{fungsi_hsh}' tidak ditemukan exact match di benchmark survei. Mencoba fuzzy matching...")
            match_found = False
            for idx, row in skor_benchmark_survei.iterrows():
                benchmark_hsh_norm = row['HSH_normalized']
                if fungsi_hsh_normalized in benchmark_hsh_norm or benchmark_hsh_norm in fungsi_hsh_normalized:
                    benchmark_data = skor_benchmark_survei.iloc[[idx]]
                    st.info(f"✓ Ditemukan match: '{row.iloc[0]}' untuk HSH '{fungsi_hsh}'")
                    match_found = True
                    break
            
            if not match_found:
                st.warning(f"⚠️ Data benchmark survei untuk HSH '{fungsi_hsh}' tidak ditemukan. Menggunakan benchmark 'Pertamina Group' sebagai referensi.")
                benchmark_data = skor_benchmark_survei[
                    skor_benchmark_survei['HSH_normalized'].str.contains('PERTAMINA GROUP', na=False)
                ]
                if benchmark_data.empty:
                    benchmark_data = skor_benchmark_survei.iloc[[0]]
                    st.info(f"Menggunakan benchmark: '{benchmark_data.iloc[0, 0]}'")
        
        # Ambil nilai dengan penanganan error yang lebih baik
        def safe_get_value(row, column_name, default='N/A'):
            try:
                if column_name in row:
                    value = row[column_name]
                    return value if pd.notna(value) else default
                return default
            except:
                return default
        
        skor_survei_val = safe_get_value(fungsi_data.iloc[0], 'Skor Survei', 'N/A')
        skor_pekerja_val = safe_get_value(fungsi_data.iloc[0], 'SKOR PEKERJA', 'N/A')
        skor_mitra_val = safe_get_value(fungsi_data.iloc[0], 'SKOR MITRA KERJA', 'N/A')
        
        p_akhlak = safe_get_value(fungsi_data.iloc[0], 'P. AKHLAK', 'N/A')
        p_one = safe_get_value(fungsi_data.iloc[0], 'P. ONE Pertamina', 'N/A')
        p_program = safe_get_value(fungsi_data.iloc[0], 'P. Program Budaya', 'N/A')
        p_keberlanjutan = safe_get_value(fungsi_data.iloc[0], 'P. Keberlanjutan', 'N/A')
        p_safety = safe_get_value(fungsi_data.iloc[0], 'P. Safety', 'N/A')
        
        mk_akhlak = safe_get_value(fungsi_data.iloc[0], 'MK. AKHLAK', 'N/A')
        mk_one = safe_get_value(fungsi_data.iloc[0], 'MK. ONE Pertamina', 'N/A')
        mk_program = safe_get_value(fungsi_data.iloc[0], 'MK. Program Budaya', 'N/A')
        mk_keberlanjutan = safe_get_value(fungsi_data.iloc[0], 'MK. Keberlanjutan', 'N/A')
        mk_safety = safe_get_value(fungsi_data.iloc[0], 'MK. Safety', 'N/A')
        
        benchmark_pekerja = benchmark_data.iloc[0, 6] if len(benchmark_data.columns) > 6 else 'N/A'
        benchmark_mitra = benchmark_data.iloc[0, 12] if len(benchmark_data.columns) > 12 else 'N/A'
        benchmark_survei = benchmark_data.iloc[0, 13] if len(benchmark_data.columns) > 13 else 'N/A'
        
        b_p_akhlak = benchmark_data.iloc[0, 1] if len(benchmark_data.columns) > 1 else 'N/A'
        b_p_one = benchmark_data.iloc[0, 2] if len(benchmark_data.columns) > 2 else 'N/A'
        b_p_program = benchmark_data.iloc[0, 3] if len(benchmark_data.columns) > 3 else 'N/A'
        b_p_keberlanjutan = benchmark_data.iloc[0, 4] if len(benchmark_data.columns) > 4 else 'N/A'
        b_p_safety = benchmark_data.iloc[0, 5] if len(benchmark_data.columns) > 5 else 'N/A'
        
        b_mk_akhlak = benchmark_data.iloc[0, 7] if len(benchmark_data.columns) > 7 else 'N/A'
        b_mk_one = benchmark_data.iloc[0, 8] if len(benchmark_data.columns) > 8 else 'N/A'
        b_mk_program = benchmark_data.iloc[0, 9] if len(benchmark_data.columns) > 9 else 'N/A'
        b_mk_keberlanjutan = benchmark_data.iloc[0, 10] if len(benchmark_data.columns) > 10 else 'N/A'
        b_mk_safety = benchmark_data.iloc[0, 11] if len(benchmark_data.columns) > 11 else 'N/A'
        
        # Hitung selisih dengan penanganan error
        def calculate_difference(val1, val2):
            try:
                if val1 == 'N/A' or val2 == 'N/A':
                    return 'N/A'
                v1 = float(str(val1).replace(',', '.'))
                v2 = float(str(val2).replace(',', '.'))
                return v1 - v2
            except:
                return 'N/A'
        
        diff_survei = calculate_difference(skor_survei_val, benchmark_survei)
        diff_pekerja = calculate_difference(skor_pekerja_val, benchmark_pekerja)
        diff_mitra = calculate_difference(skor_mitra_val, benchmark_mitra)
        
        benchmark_hsh_display = benchmark_data.iloc[0, 0] if not benchmark_data.empty else "Benchmark tidak tersedia"
        
        comparison_text = f"""
PERBANDINGAN SKOR SURVEI

Fungsi: {selected_fungsi}
HSH Fungsi: {fungsi_hsh}
HSH Benchmark: {benchmark_hsh_display}

=== RINGKASAN SKOR FUNGSI ===
• Skor Survei Total: {skor_survei_val}
• SKOR PEKERJA: {skor_pekerja_val}
  - P. AKHLAK: {p_akhlak}
  - P. ONE Pertamina: {p_one}
  - P. Program Budaya: {p_program}
  - P. Keberlanjutan: {p_keberlanjutan}
  - P. Safety: {p_safety}

• SKOR MITRA KERJA: {skor_mitra_val}
  - MK. AKHLAK: {mk_akhlak}
  - MK. ONE Pertamina: {mk_one}
  - MK. Program Budaya: {mk_program}
  - MK. Keberlanjutan: {mk_keberlanjutan}
  - MK. Safety: {mk_safety}

=== BENCHMARK ({benchmark_hsh_display}) ===
• Skor Survei Total: {benchmark_survei}
• SKOR PEKERJA: {benchmark_pekerja}
  - P. AKHLAK: {b_p_akhlak}
  - P. ONE Pertamina: {b_p_one}
  - P. Program Budaya: {b_p_program}
  - P. Keberlanjutan: {b_p_keberlanjutan}
  - P. Safety: {b_p_safety}

• SKOR MITRA KERJA: {benchmark_mitra}
  - MK. AKHLAK: {b_mk_akhlak}
  - MK. ONE Pertamina: {b_mk_one}
  - MK. Program Budaya: {b_mk_program}
  - MK. Keberlanjutan: {b_mk_keberlanjutan}
  - MK. Safety: {b_mk_safety}

=== SELISIH (Fungsi - Benchmark) ===
• Skor Survei Total: {diff_survei} {'✓' if diff_survei != 'N/A' and isinstance(diff_survei, (int, float)) and diff_survei > 0 else '⚠' if diff_survei != 'N/A' and isinstance(diff_survei, (int, float)) and diff_survei < 0 else ''}
• SKOR PEKERJA: {diff_pekerja} {'✓' if diff_pekerja != 'N/A' and isinstance(diff_pekerja, (int, float)) and diff_pekerja > 0 else '⚠' if diff_pekerja != 'N/A' and isinstance(diff_pekerja, (int, float)) and diff_pekerja < 0 else ''}
• SKOR MITRA KERJA: {diff_mitra} {'✓' if diff_mitra != 'N/A' and isinstance(diff_mitra, (int, float)) and diff_mitra > 0 else '⚠' if diff_mitra != 'N/A' and isinstance(diff_mitra, (int, float)) and diff_mitra < 0 else ''}

Catatan:
✓ = Fungsi LEBIH BAIK dari benchmark
⚠ = Fungsi memiliki PELUANG PENGEMBANGAN
"""
        
        prompt = f"""
**KONTEKS:** Analisis perbandingan Survei untuk fungsi {selected_fungsi} di HSH {selected_hsh}

**DATA SURVEI YANG DISEDIAKAN:**
{comparison_text}

**INSTRUKSI ANALISIS:**
1. ANALISIS persepsi pekerja dan mitra kerja terhadap implementasi budaya
2. IDENTIFIKASI pola dan area yang memerlukan perhatian khusus
3. BERIKAN REASONING untuk setiap temuan dan rekomendasi
4. FOKUS pada aspek perilaku yang mempengaruhi persepsi

**ASPEK YANG DINILAI:**
- Pemahaman dan penerapan nilai AKHLAK
- Implementasi ONE Pertamina
- Partisipasi dalam Program Budaya
- Komitmen terhadap Keberlanjutan
- Budaya Safety

**FORMAT OUTPUT YANG DIHARAPKAN:**

**Apresiasi Pencapaian:**
[Apresiasi terhadap skor yang sudah di atas atau sesuai benchmark, soroti area kekuatan dalam persepsi]

**Hal yang Sudah Baik:**
- [Area spesifik 1] - **Reasoning:** [Penjelasan mengapa area ini mendapat persepsi positif]
- [Area spesifik 2] - **Reasoning:** [Penjelasan mengapa area ini mendapat persepsi positif]

**Peluang Pengembangan Lebih Lanjut:**
- [Area 1] - **Reasoning:** [Penjelasan logis dan saran konkret untuk meningkatkan persepsi]
- [Area 2] - **Reasoning:** [Penjelasan logis dan saran konkret untuk meningkatkan persepsi]

**Strategi Perbaikan Persepsi:**
[Berikan 1-2 strategi spesifik untuk meningkatkan persepsi pekerja dan mitra kerja]
"""
        return call_deepseek(prompt, max_tokens=1500)
    except Exception as e:
        error_msg = f"Error dalam analisis survei: {str(e)}\n\nDetail error: {e.__class__.__name__}"
        st.error(error_msg)
        return error_msg

def create_word_document(fungsi_name, analyses):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    title = doc.add_heading('Rapport Writer Assistance', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    today = datetime.now().strftime('%d %B %Y')
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f'Laporan Analisis Implementasi Budaya Kerja\n{fungsi_name}\n{today}')
    subtitle_run.bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    doc.add_paragraph()
    
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        'Laporan ini disusun dengan pendekatan apresiatif untuk memberikan gambaran komprehensif '
        'mengenai implementasi budaya kerja dengan fokus pada aspek perilaku (behavior). '
        'Analisis dilakukan berdasarkan data evidence, survei, dan perbandingan dengan benchmark.'
    )
    intro_run.italic = True
    doc.add_paragraph()
    
    doc.add_heading('1. Analisis Strategi Budaya', 1)
    doc.add_paragraph(analyses['strategi_budaya'])
    doc.add_paragraph()
    
    doc.add_heading('2. Analisis Program Budaya', 1)
    doc.add_paragraph(analyses['program_budaya'])
    doc.add_paragraph()
    
    doc.add_heading('3. Analisis Impact to Business', 1)
    doc.add_paragraph(analyses['impact'])
    doc.add_paragraph()
    
    doc.add_heading('4. Analisis Perbandingan Evidence dengan Benchmark', 1)
    doc.add_paragraph(analyses['evidence_comparison'])
    doc.add_paragraph()
    
    doc.add_heading('5. Analisis Perbandingan Survei dengan Benchmark', 1)
    doc.add_paragraph(analyses['survei_comparison'])
    doc.add_paragraph()
    
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    doc.add_paragraph()
    
    closing = doc.add_paragraph()
    closing_run = closing.add_run(
        'Laporan ini disusun sebagai bahan refleksi dan pengembangan berkelanjutan dalam implementasi '
        'budaya kerja. Kami mengapresiasi komitmen dan dedikasi seluruh tim dalam mewujudkan '
        'transformasi budaya yang positif dan berkelanjutan.'
    )
    closing_run.italic = True
    
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(f'\nDibuat oleh Rapport Writer Assistance\n{datetime.now().strftime("%d %B %Y, %H:%M WIB")}').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# Main App
def main():
    st.title("📊 Rapport Writer Assistance")
    st.caption("Asisten Analisis Implementasi Budaya Kerja dengan Pendekatan Apresiatif")
    
    with st.expander("📖 PETUNJUK PENGGUNAAN", expanded=True):
        st.markdown("""
        ### Selamat Datang di Rapport Writer Assistance!
        
        Aplikasi ini menggunakan **DeepSeek AI** dan tampilan **nuansa coklat** yang hangat.
        
        **Langkah-langkah Penggunaan:**
        1. Pilih HSH dan Fungsi di sidebar
        2. Upload file PCB dan Impact (opsional)
        3. Klik **"🚀 Mulai Analisis"**
        4. Download hasil dalam format .docx
      
        """)

    with st.spinner('Memuat data...'):
        skor_total, skor_survei, skor_benchmark_evidence, skor_benchmark_survei = load_excel_files()
    
    if skor_total is None:
        st.stop()
    
    st.sidebar.header("⚙️ Pengaturan Analisis")
    
    hsh_list = sorted(skor_total['HSH'].unique().tolist())
    selected_hsh = st.sidebar.selectbox("Pilih HSH:", options=hsh_list)
    filtered_fungsi = sorted(skor_total[skor_total['HSH'] == selected_hsh]['Fungsi'].unique().tolist())
    selected_fungsi = st.sidebar.selectbox("Pilih Fungsi:", options=filtered_fungsi)
    
    st.sidebar.markdown("---")
    st.sidebar.subheader("📁 Upload Dokumen")
    uploaded_pcb = st.sidebar.file_uploader("Upload PCB", type=['xlsx', 'xls', 'pdf', 'png', 'jpg', 'jpeg'])
    uploaded_impact = st.sidebar.file_uploader("Upload Impact to Business", type=['xlsx', 'xls', 'pdf', 'png', 'jpg', 'jpeg'])
    st.sidebar.markdown("---")
    
    # 🟤 TOMBOL MULAI ANALISIS - NUANSA COKLAT
    st.markdown("""
    <style>
    div.stButton > button {
        background-color: #5d4037;
        color: white;
        border: none;
        padding: 12px 24px;
        border-radius: 8px;
        font-weight: bold;
        font-size: 16px;
        transition: background-color 0.3s ease;
        width: 100%;
    }
    div.stButton > button:hover {
        background-color: #4e342e;
    }
    </style>
    """, unsafe_allow_html=True)

    analyze_button = st.sidebar.button("🚀 Mulai Analisis", use_container_width=True)
    
    if analyze_button:
        if uploaded_pcb is None:
            st.error("⚠️ Silakan upload file PCB terlebih dahulu!")
            st.stop()
        
        st.success(f"✅ Memproses analisis untuk **{selected_fungsi}** (HSH: {selected_hsh})")
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        pcb_content = read_uploaded_file(uploaded_pcb)
        impact_content = read_uploaded_file(uploaded_impact) if uploaded_impact else None
        
        status_text.text("🔍 Menganalisis Strategi Budaya...")
        progress_bar.progress(25)
        strategi_budaya = analyze_strategi_budaya(pcb_content, selected_hsh, selected_fungsi)
        
        status_text.text("🔍 Menganalisis Program Budaya...")
        progress_bar.progress(40)
        program_budaya = analyze_program_budaya(pcb_content, selected_hsh, selected_fungsi)
        
        status_text.text("🔍 Menganalisis Impact to Business...")
        progress_bar.progress(55)
        impact = analyze_impact(impact_content, selected_hsh, selected_fungsi)
        
        status_text.text("🔍 Menganalisis Perbandingan Evidence...")
        progress_bar.progress(70)
        evidence_comparison = analyze_evidence_comparison(skor_total, skor_benchmark_evidence, selected_hsh, selected_fungsi)
        
        status_text.text("🔍 Menganalisis Perbandingan Survei...")
        progress_bar.progress(85)
        survei_comparison = analyze_survei_comparison(skor_survei, skor_benchmark_survei, selected_hsh, selected_fungsi)
        
        analyses = {
            'strategi_budaya': strategi_budaya,
            'program_budaya': program_budaya,
            'impact': impact,
            'evidence_comparison': evidence_comparison,
            'survei_comparison': survei_comparison
        }
        
        status_text.text("📝 Membuat dokumen Word...")
        progress_bar.progress(95)
        doc_io = create_word_document(selected_fungsi, analyses)
        
        progress_bar.progress(100)
        status_text.text("✅ Analisis selesai!")
        st.balloons()
        
        st.markdown("---")
        st.header("📊 Hasil Analisis")
        
        # 🟤 TAB - NUANSA COKLAT
        st.markdown("""
        <style>
        .stTabs [data-baseweb="tab-list"] {
            background-color: #f5f0e6;
            padding: 10px;
            border-radius: 8px;
        }
        .stTabs [data-baseweb="tab"] {
            height: 40px;
            white-space: pre-wrap;
            background-color: #e8dccf;
            border-radius: 6px;
            color: #4e342e;
            font-weight: bold;
            padding: 0 16px;
            margin-right: 8px;
        }
        .stTabs [aria-selected="true"] {
            background-color: #5d4037;
            color: white;
        }
        </style>
        """, unsafe_allow_html=True)

        tab1, tab2, tab3, tab4, tab5 = st.tabs([
            "Strategi Budaya", 
            "Program Budaya", 
            "Impact to Business",
            "Perbandingan Evidence",
            "Perbandingan Survei"
        ])
        
        with tab1: st.markdown("### Analisis Strategi Budaya\n" + strategi_budaya)
        with tab2: st.markdown("### Analisis Program Budaya\n" + program_budaya)
        with tab3: st.markdown("### Analisis Impact to Business\n" + impact)
        with tab4: st.markdown("### Analisis Perbandingan Evidence\n" + evidence_comparison)
        with tab5: st.markdown("### Analisis Perbandingan Survei\n" + survei_comparison)
        
        st.markdown("---")
        today = datetime.now().strftime('%m_%d')
        filename = f"Rapp_{selected_fungsi.replace(' ', '_').replace('/', '_')}_{today}.docx"

        # 🟤 TOMBOL DOWNLOAD - NUANSA COKLAT
        st.markdown("""
        <style>
        .stDownloadButton > button {
            background-color: #5d4037 !important;
            color: white !important;
            border: none !important;
            padding: 12px 24px !important;
            border-radius: 8px !important;
            font-weight: bold !important;
            font-size: 16px !important;
            width: 100% !important;
            transition: background-color 0.3s ease !important;
        }
        .stDownloadButton > button:hover {
            background-color: #4e342e !important;
        }
        </style>
        """, unsafe_allow_html=True)

        st.download_button(
            label="📥 Download Hasil Analisis (.docx)",
            data=doc_io,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

        # 🟤 PESAN SUKSES - NUANSA COKLAT MUDA
        st.markdown(f"""
        <div style="
            background-color: #f9f4ed;
            padding: 12px;
            border-radius: 8px;
            border-left: 4px solid #5d4037;
            margin-top: 10px;
            color: #4e342e;
            font-weight: bold;
        ">
            ✅ Dokumen siap didownload: <strong>{filename}</strong>
        </div>
        """, unsafe_allow_html=True)
    
    else:
        st.info("👈 Silakan pilih HSH, Fungsi, upload file, dan klik tombol **Mulai Analisis** di sidebar")
        col1, col2 = st.columns(2)
        with col1: st.metric("HSH Terpilih", selected_hsh if selected_hsh else "-")
        with col2: st.metric("Fungsi Terpilih", selected_fungsi if selected_fungsi else "-")
        
        st.markdown("---")
        st.markdown("### 💡 Tips Penggunaan")
        st.markdown("""
        - Semua analisis menggunakan **pendekatan apresiatif**
        - Fokus pada **perubahan perilaku**, bukan teknis
        - API key disimpan aman melalui **Streamlit Secrets**
        - Hasil analisis di-cache untuk **penghematan biaya** dan kecepatan
        """)

if __name__ == "__main__":
    main()
